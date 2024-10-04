
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def create_heading(doc, text, level=1):
    """Creates a heading in the Word document. Level 1 headings will be size 16, level 2 will be size 14."""
    doc.add_paragraph()  # Add an empty line before the heading for better spacing
    heading = doc.add_paragraph()
    run = heading.add_run(text)
    if level == 0:
        run.font.size = Pt(20) # Main heading size 20
    elif level == 1:
        run.font.size = Pt(16)  # Main heading size 16
    else:
        run.font.size = Pt(14)  # Subheading size 14

    run.bold = True  # Make it bold
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
# Function to create content paragraphs with size 12 and handle bold text
def create_content_paragraph(doc, text):
    """Creates a content paragraph with size 12 Arial font and bolds text enclosed in **."""
    paragraph = doc.add_paragraph()  # Create a new paragraph
    # Split by double asterisks to identify bold segments
    parts = re.split(r"(\*\*.*?\*\*)", text.strip())  # Ensure no trailing/leading spaces

    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # Remove the asterisks and make this part bold
            part = part[2:-2].strip()
            run = paragraph.add_run(part)
            run.bold = True
        else:
            run = paragraph.add_run(part)

        run.font.size = Pt(12)
        run.font.name = 'Arial'

# Function to create numbered headings with bold formatting
def create_numbered_heading(doc, text, num, level=1):
    """Creates a numbered bold heading in the Word document"""
    heading = doc.add_paragraph()
    run = heading.add_run(f'{num} {text}')
    run.font.name = 'Arial'
    run.font.size = Pt(12)  # All text is set to size 12
    run.bold = True  # Make the heading bold
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

# Function to create paragraphs with size 12, normal text
def create_paragraph(doc, text):
    """Creates a normal paragraph with size 12 Arial font"""
    paragraph = doc.add_paragraph(text)
    run = paragraph.runs[0]
    run.font.size = Pt(12)  # Standard font size for description
    run.font.name = 'Arial'
    run.bold = False  # Ensure the paragraphs are not bold

# Function to create numbered bullet headings with size 16, bold formatting
def create_bullet_numbered_heading(doc, text, num):
    """Creates a numbered bullet heading in the Word document with bold and size 16 font"""
    heading = doc.add_paragraph(style='List Number')  # Use 'List Number' for numbered bullets
    run = heading.add_run(f'{text}')  # No need to manually add the number, Word handles it
    run.font.name = 'Arial'
    run.font.size = Pt(16)  # Heading size 16
    run.bold = True  # Make it bold
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

# Function to create content paragraphs with bullet points and size 14
def create_bullet_content_paragraph(doc, text):
    """Creates a content paragraph with bullet points and size 14 Arial font"""
    paragraph = doc.add_paragraph(style='List Bullet')  # Use 'List Bullet' for bullet points
    run = paragraph.add_run(text)
    run.font.size = Pt(14)  # Content font size 14
    run.font.name = 'Arial'

    
# Function to process content into numbered bullet headings and bullet point paragraphs
def process_content(doc, content):
    """Processes the content and formats it into the Word document"""
    lines = content.strip().split("\n")  # Strip to remove leading/trailing whitespace

    for line in lines:
        line = line.strip()  # Remove extra spaces
        if line.startswith("##"):
            create_heading(doc, line[2:],level=0)
            continue
        # Check for main headings (e.g., **1. Heading**)
        if re.match(r"\*\*\d+\..*\*\*", line):
            heading_text = line[2:-2].strip()  # Remove the ** around the heading
            create_heading(doc, heading_text, level=1)

        # Check for subheadings (e.g., **1.1 Sub Heading**)
        elif re.match(r"\*\*\d+\.\d+.*\*\*", line):
            subheading_text = line[2:-2].strip()  # Remove the ** around the subheading
            create_heading(doc, subheading_text, level=2)

        # For other lines (content)
        elif line:
            line = line.replace('* ', '').replace("**", '').replace("*",'').strip()  # Remove leading bullet points if any
            create_content_paragraph(doc, line)

# Function to generate Word document from uploaded txt file
def generate_word_from_txt(content , File_Name ):
    # Read the content from the uploaded text file
    
    # Create a new Word document
    doc = Document()

    # Set default document style to Arial and size 12
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)

    # Process the content and format it into the document
    process_content(doc, content)

    # Save the document as a Word file
    output_file = File_Name.replace(".txt", ".docx")
    doc.save(os.path.join(os.getcwd(),"AllDocuments" ,output_file))
    print(f"Document saved as {output_file}")
    return output_file
def txt_to_docx(txt_file, docx_file):
    # Create a new Document object
    doc = Document()

    # Regular expression to match a sentence starting with a number
    num_start_pattern = re.compile(r'^\d+\.?\s*')

    # Open the txt file and read its content
    with open(txt_file, 'r') as file:
        lines = file.readlines()

    # Add each line from the txt file to the docx document
    for line in lines:
        # Check if the line starts with a number
        match = num_start_pattern.match(line)
        if match:
            # Create a paragraph
            paragraph = doc.add_paragraph()
            
            # Get the matched number and text
            number = match.group()
            rest_of_line = line[len(number):]

            # Add bold text for the number part
            paragraph.add_run(number).bold = True

            # Add the rest of the line
            paragraph.add_run(rest_of_line)
        else:
            # If no number at the start, just add the full line
            doc.add_paragraph(line)

    # Save the docx file
    output_file = txt_file.replace(".txt", ".docx")
    doc.save(os.path.join(os.getcwd(),"AllDocuments" ,output_file))
    print(f"Successfully converted '{txt_file}' to '{docx_file}'.")
def text_doc(textString , txt_file):
    
    doc = Document()
    
    output_file = txt_file.replace(".txt", ".docx")
    doc.add_paragraph(textString)
    
    
    doc.save(os.path.join(os.getcwd(),"AllDocuments" ,output_file))
from docx import Document

def parse_input_string(input_string):
    # Split the input string by lines and then by the delimiter '|'
    lines = input_string.strip().split('\n')
    
    # Remove empty lines and strip excess whitespace
    parsed_data = []
    for line in lines:
        if line.strip():  # Ignore empty lines
            # Split the line based on '|' and strip extra spaces from each cell
            parsed_data.append([cell.strip() for cell in line.split('|') if cell.strip()])
    
    return parsed_data

def create_table_doc(data, txt_file):
    # Create a new Document
    doc = Document()
    doc.add_heading('Component Overview Table', 0)
    print("Came here")
    # Add the table with appropriate number of rows and columns
    num_rows = len(data)
    num_cols = len(data[0])

    # Define table
    table = doc.add_table(rows=num_rows, cols=num_cols)

    # Style the table
    table.style = 'Table Grid'

    # Insert the data into the table
    for i, row in enumerate(data):
        for j, cell_text in enumerate(row):
            cell = table.cell(i, j)
            cell.text = cell_text
            if i == 0:  # Header row
                cell_font = cell.paragraphs[0].runs[0].font
                cell_font.bold = True

    # Save the document
    output_file = txt_file.replace(".txt", ".docx")
    doc.save(os.path.join(os.getcwd(),"AllDocuments" ,output_file))

# Example input in string format
def TOLS_text_doc(data , txt_file):
    
    # Create a new .docx document
    doc = Document()

    # Add a title to the document
    doc.add_heading('Table of Limitations (TOL)', 0)

    # Add a brief description
    doc.add_paragraph('This document contains the Table of Limitations based on the LLD provided.')

    # Split the generated TOL text into lines (for parsing)
    tol_lines = data.split('\n')

    # Start adding the table (3 columns)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'

    # Add header row to the table
    hdr_cells = table.rows[0].cells

    listNew = [' Test Operation ID',' Operation Description',' Linked Requirements','---']
    # Add rows from TOL lines
    for line in tol_lines:
        # Assuming the generated text follows a consistent format, parse each line
        if line.strip():  # Check if the line is not empty
            # Split the line into columns based on tabs or another separator
            columns = line.split('|')  # Adjust separator if needed 
            print(columns)
            # Add a new row to the table
            row_cells = table.add_row().cells
            if len(columns)>1 and columns[1].replace(" ","") in listNew :
                continue 
            row_cells[0].text = columns[1] if len(columns) > 1 else ""
            row_cells[1].text = columns[2] if len(columns) > 2 else ""
            row_cells[2].text = columns[3] if len(columns) > 3 else ""

    # Save the document to disk
    save_path = os.path.join(os.getcwd(),"AllDocuments" ,txt_file.replace(".txt",".docx"))  # Change this to your desired path

    # Save the document to the specified path
    doc.save(save_path)

def TOLS_text_TestCase(data , txt_file):
    
    
    # Create a new .docx document
    doc = Document()

    # Add a title to the document
    doc.add_heading('Table of Test Case ', 0)

    # Add a brief description


    # Split the generated TOL text into lines (for parsing)
    tol_lines = data.split('\n')

    # Start adding the table (3 columns)
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'

    # Add header row to the table
    hdr_cells = table.rows[0].cells

    listNew = [' Test Operation ID',' Operation Description',' Linked Requirements','---']
    # Add rows from TOL lines
    for line in tol_lines:
        # Assuming the generated text follows a consistent format, parse each line
        if line.strip():  # Check if the line is not empty
            # Split the line into columns based on tabs or another separator
            columns = line.split('|')  # Adjust separator if needed 
            print(columns)
            # Add a new row to the table
            row_cells = table.add_row().cells
            if len(columns)>1 and columns[1].replace(" ","") in listNew :
                continue 
            row_cells[0].text = columns[1] if len(columns) > 1 else ""
            row_cells[1].text = columns[2] if len(columns) > 2 else ""
            row_cells[2].text = columns[3] if len(columns) > 3 else ""
            row_cells[3].text = columns[4] if len(columns) > 4 else ""
            row_cells[4].text = columns[5] if len(columns) > 5 else ""
            

    
    save_path = os.path.join(os.getcwd(),"AllDocuments" ,txt_file.replace(".txt",".docx"))  # Change this to your desired path

    # Save the document to the specified path
    doc.save(save_path)


if __name__ =="__main__":
    txt_to_docx()

    # Parse the string input
    parsed_data = parse_input_string(input_string)

    # Specify the file name
    output_file = "component_overview_table.docx"

    # Call the function to create the Word document
    create_table_doc(parsed_data, output_file)

    print(f"Document created: {output_file}")